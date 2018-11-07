import string
import random
import json
from django.contrib.auth.hashers import make_password
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.serializers import json
from django.http import HttpResponseRedirect, HttpResponse,HttpResponseForbidden,JsonResponse
from django.shortcuts import render, redirect
from django.core.mail import send_mail
# from django.core.mail import EmailMultiAlternatives
from django.http import FileResponse
from django.utils import timezone
from reportlab.pdfgen import canvas
from shiyanshi.settings import EMAIL_FROM
import docx

# Create your views here.
# 基于类实现需要继承的view
from django.urls import reverse
from django.views.generic.base import View
from users.forms import LoginForm, UserInfoForm, UploadImageForm, \
    ModifyPwdForm, RegisterForm,SeatForm, RegForm,ForgotPasswordForm
# Django自带的用户验证,login
from django.contrib.auth import authenticate, login, logout
from users.models import Seat, UserProfile, Banner, Week, WTime, User_ex, EmailVerifyRecord
from utils.email_send import send_register_eamil

# 登录
class LoginView(View):
    # 直接调用get方法免去判断
    def get(self, request):
        # render就是渲染html返回用户
        # render三变量: request 模板名称 一个字典写明传给前端的值
        redirect_url = request.GET.get('next', '')
        return render(request, "login.html", {
            "redirect_url": redirect_url
        })

    def post(self, request):
        # 类实例化需要一个字典参数dict:request.POST就是一个QueryDict所以直接传入
        # POST中的usernamepassword，会对应到form中
        login_form = LoginForm(request.POST)

        # is_valid判断我们字段是否有错执行我们原有逻辑，验证失败跳回login页面
        if login_form.is_valid():
            # 取不到时为空，username，password为前端页面name值
            # user_name = login_form.cleaned_data['username']
            # pass_word = login_form.cleaned_data['password']
            user_name = request.POST.get("username", "")
            pass_word = request.POST.get("password", "")
            # 成功返回user对象,失败返回null
            user = authenticate(username=user_name, password=pass_word)

            # 如果不是null说明验证成功
            if user is not None:
                # 只有当用户激活时才给登录
                if user.is_active:
                    login(request, user)

                    return HttpResponseRedirect(reverse("index"))
                # 即用户未激活跳转登录，提示未激活
                else:
                    return render(
                        request, "login.html", {
                            "msg": "用户名未激活! 请前往邮箱进行激活"})
            # 仅当用户真的密码出错时
            else:
                return render(request, "login.html", {"msg": "用户名或密码错误!"})
        # 验证不成功跳回登录页面
        # 没有成功说明里面的值是None，并再次跳转回主页面
        else:
            return render(
                request, "login.html", {
                    "login_form": login_form})


class LogoutView(View):
    def get(self, request):
        # django自带的logout
        logout(request)
        return render(request, 'login.html', {})
        # 重定向到首页
        # return HttpResponseRedirect(reverse("index"))


# 首页view
class IndexView(LoginRequiredMixin, View):
    def get(self, request):
        all_banner = Banner.objects.all()
        return render(request, 'index.html', {'all_banner': all_banner})


class ChoiceTime(LoginRequiredMixin,View):
    def get(self, request):
        b = []
        f = ""
        # 获取时间 星期 节课
        date = request.GET.get('date')
        time = request.GET.get('time')
        # 获取select的值
        c = request.GET.get('index')
        d = request.GET.get('index1')

        """错误处理"""
        try:
            date1 = int(date)
            time1 = int(time)
            c1 = int(c)
            d1 = int(d)
            if date1 < 0 or date1 > 7 :
                print("data")
                return HttpResponseForbidden
            if time1 != 12 and time1 != 34 and time1 != 56 and time1 != 78 and time1 != 90:
                print("timr")
                return HttpResponseForbidden
            if c1<0 or c1>7 :
                print("c1")
                return HttpResponseForbidden
            if d1 < 0 or d1 > 5:
                print("c2")
                return HttpResponseForbidden
        except Exception as e:
            return HttpResponseForbidden

        weeks = Week.objects.filter(dateTime=date)

        wtimes = WTime.objects.filter(week__in=weeks, classtime=time)

        seats = Seat.objects.filter(week__in=wtimes)

        if seats:
            for seat in seats:
                status = seat.status
                if status == '0':
                    print(seat.number)
                    i, j = train(seat.number)
                    f = (str(i) + "_" + str(j))
                    b.append(f)

        return render(request, 'getSeat.html', {'c': c, 'b2': b, 'd': d})

    def post(self, request):
        b = []
        f =" "
        return HttpResponse('{"status":"success", "msg":"你已经有预约了"}', content_type='application/json')



class Seats(LoginRequiredMixin,View):
    """
    初始化界面 把正在实验 status状态为0的数 取出来
    """

    def get(self, request):
        b = []
        return render(request, 'getSeat.html', {'b2': b, 'c': 0, 'd': 0})

    def post(self, request, ab=None, all_seat=None):
        user_name = request.user
        label = request.POST.get('label')
        date = request.POST.get('date')
        time = request.POST.get('time')
        # label = tramsport(label)
        # print(label,date,time)

        # 设置座位不可选状态的返回 错误处理 如果标签纸为0 返回刷新页面
        # if label == "0":
        #     return HttpResponse('{"status":"fail"}', content_type='application/json')

        # 设置座位不可选状态的返回 错误处理 如果标签纸为0 返回刷新页面
        if label == "0":
            return HttpResponse('{"status":"fail"}', content_type='application/json')
        label = tramsport(label)
        """查询当前用户是否已经选座"""
        current_user = UserProfile.objects.filter(username=user_name)
        for users in current_user:
            users = users
        if users.seat_status == "0":
            return HttpResponse('{"status":"fail", "msg":"你已经有预约了"}', content_type='application/json')

        """根据座位号查询该座位信息"""
        weeks = Week.objects.filter(dateTime=date)

        wtimes = WTime.objects.filter(week__in=weeks, classtime=time)
        print(wtimes)
        seats = Seat.objects.filter(week__in=wtimes, number=label)
        print(seats)
        for seat in seats:
            current_seat_status = seat.status
            if current_seat_status == "0":
                return HttpResponse('{"status":"fail", "msg":"你已经有预约了"}', content_type='application/json')
            if current_seat_status == "1":
                seat.status = "0"
                seat.user = users
                users.seat_status = "0"
                users.save()
                seat.save()
        return HttpResponse('{"status":"success", "msg":"收藏"}', content_type='application/json')


def train(number):
    i = (number+2) // 3
    j = number % 3
    if i == 0:
        i = 1
    if j == 0:
        j=3
    return i,j


def tramsport(label):
    a = label.split('_')[0]
    b = label.split('_')[1]
    a = int(a)
    b = int(b)
    c = a*3-(3-b)
    return c


class UserInfoView(LoginRequiredMixin,View):
    # login_url = '/login/'
    # redirect_field_name = 'next'

    def get(self, request):
        return render(request, "usercenter-info.html", {

        })

    def post(self, request):
        # 不像用户咨询是一个新的。需要指明instance。不然无法修改，而是新增用户
        user_info_form = UserInfoForm(request.POST, instance=request.user)
        if user_info_form.is_valid():
            user_info_form.save()
            return HttpResponse(
                '{"status":"success"}',
                content_type='application/json')
        else:
            # 通过json的dumps方法把字典转换为json字符串
            return HttpResponse(
                json.dumps(
                    user_info_form.errors),
                content_type='application/json')


class UploadImageView( View):
    # login_url = '/login/'
    # redirect_field_name = 'next'

    def post(self, request):
        # 这时候用户上传的文件就已经被保存到imageform了 ，为modelform添加instance值直接保存
        image_form = UploadImageForm(
            request.POST, request.FILES, instance=request.user)
        if image_form.is_valid():
            image_form.save()
            # # 取出cleaned data中的值,一个dict
            # image = image_form.cleaned_data['image']
            # request.user.image = image
            # request.user.save()
            return HttpResponse(
                '{"status":"success"}',
                content_type='application/json')
        else:
            return HttpResponse(
                '{"status":"fail"}',
                content_type='application/json')


class UpdatePwdView(LoginRequiredMixin, View):
    # login_url = '/login/'
    # redirect_field_name = 'next'

    def post(self, request):
        modify_form = ModifyPwdForm(request.POST)
        if modify_form.is_valid():
            pwd1 = request.POST.get("password1", "")
            pwd2 = request.POST.get("password2", "")
            # 如果两次密码不相等，返回错误信息
            if pwd1 != pwd2:
                return HttpResponse(
                    '{"status":"fail", "msg":"密码不一致"}',
                    content_type='application/json')
            # 如果密码一致
            user = request.user
            # 加密成密文
            user.password = make_password(pwd2)
            # save保存到数据库
            user.save()
            return HttpResponse(
                '{"status":"success"}',
                content_type='application/json')
        # 验证失败说明密码位数不够。
        else:
            # 通过json的dumps方法把字典转换为json字符串
            return HttpResponse(
                json.dumps(
                    modify_form.errors),
                content_type='application/json')


class BespeakView(LoginRequiredMixin,View):
    def get(self, request):
        # 获取文档对象
        result= []
        file = docx.Document("F:\pyy\shiyanshi\media/files\实验室管理制度.docx")
        print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

        # 输出每一段的内容
        for para in file.paragraphs:
            print(para.text)
            result.append(para.text)

        # 输出段落编号及段落内容
        for i in range(len(file.paragraphs)):
            print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)

        # name = request.user.names
        # user = UserProfile.objects.get(names=name)
        # if user:
        #     word = user.pdf_url


        # word = ""
        # f = open("E:\py\Laboratory\前端开发者手册.pdf", "rb")
        # for i in f:
        #     word += i
        #     print(word)

        return render(request, 'files.html', {'result': result})


def page_not_found(request):
    # 全局404处理函数
    from django.shortcuts import render_to_response
    response = render_to_response('404.html', {})
    response.status_code = 404
    return response


def page_error(request):
    # 全局500处理函数
    from django.shortcuts import render_to_response
    response = render_to_response('500.html', {})
    response.status_code = 500
    return response


# 注册功能的view
class RegisterView(View):
    # get方法直接返回页面
    def get(self, request):
        reg_form = RegForm()
        return render(
            request, "register.html", {
                'register_form': reg_form})

    def post(self, request):

        reg_form = RegForm(request.POST)
        if reg_form.is_valid():

            user_name = request.POST.get('username')
            email = request.POST.get('email')
            pass_word = request.POST.get('password')

            # 查看用户是否注册

            if UserProfile.objects.filter(username=user_name):
                print("gggggggggg")
                return render(
                    request, "register.html", {
                        # "register_form": register_form, "msg": "用户已存在"})
                            "register_form": reg_form, "msg": "用户已存在"})
            # 查看邮箱是否注册
            email = request.POST.get('email')
            print(email)
            email_exsit = UserProfile.objects.filter(email=email)
            if email_exsit:
                return render(
                    request, "register.html", {
                        "register_form": reg_form, "msg": "该邮箱已存在存在"})

            # 查看验证码是否为空
            # verification_code = request.POST.get('verification_code', '')
            # if not verification_code:
            #     return render(
            #         request, "register.html", {
            #             "register_form": reg_form, "msg3": "验证码不能为空"})

            # 查看验证码是否正确
            verification_code = request.POST.get('verification_code', '')
            user_ex = User_ex.objects.get(valid_code=verification_code)
            if user_ex.email != email:
                return render(
                    request, "register.html", {
                        "register_form": reg_form, "msg2": "验证码错误"})
            # 创建用户
            user = UserProfile.objects._create_user(username=user_name, email=email, password=pass_word)
            # 实例化一个user_profile对象，将前台值存入
            user.save()
            user_profile = UserProfile()
            # user_profile.username = user_name
            # user_profile.email = email

            # 激活用户
            user_profile.is_active = True

            # 加密password进行保存
            user_profile.password = make_password(pass_word)
            user_profile.save()

            # 跳转到登录页面
            return render(request, "login.html", )
        # 注册邮箱form验证失败
        else:
            return render(
                request, "register.html", {
                    # "register_form": register_form})
                    "register_form": reg_form})


class Word(View):
    def get(self,request):
        return render(request,'fi.html',{})


class SomeView(View):
    def get(self, request):
        # Create the HttpResponse object with the appropriate PDF headers.
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="somefilename.pdf"'

        # Create the PDF object, using the response object as its "file."
        p = canvas.Canvas(response)

        # Draw things on the PDF. Here's where the PDF generation happens.
        # See the ReportLab documentation for the full list of functionality.
        p.drawString(100, 100, "Hello world.")

        # Close the PDF object cleanly, and we're done.
        p.showPage()
        p.save()
        return FileResponse(response)


def forgetpassword(request):
    email = request.GET.get('email', '')
    code = ''.join(random.sample(string.digits + string.ascii_letters, 6))
    print(code)
    data = {}
    send_mail(
        # subject='24小时无人'
        '24小时无人值守实验室验证码',
        '您好，你注册的验证码为：%s' % code,
        EMAIL_FROM,
        [email],
        fail_silently=False,
    )
    user_ex = User_ex.objects.create(email=email, valid_code=code, valid_time=timezone.now())
    user_ex.save()
    try:
        user_ex = User_ex.objects.filter(email=email)
        print(user_ex)

        if user_ex:
            user_ex.valid_code = code
            user_ex.save()
        else:
            print("ddddddddddddddddddddddd")
            data['message'] = u'此邮箱未注册'
            raise Exception(data['message'])
            return HttpResponse(data)
        data['status'] = 'SUCCESS'

        data['success'] = True
        data['message'] = 'OK'
        return HttpResponse(data)
    except Exception as e:
        pass


# 发送验证码
def send_verification_code(request):

    email = request.GET.get('email', '')
    code = ''.join(random.sample(string.digits + string.ascii_letters, 6))
    print(code)
    print(email)
    data = {}
    send_mail(
        # subject='24小时无人'
        '24小时无人值守实验室验证码',
        '您好，你注册的验证码为：%s' % code,
        EMAIL_FROM,
        [email],
        fail_silently=False,
    )
    user_ex = User_ex.objects.create(email=email, valid_code=code, valid_time=timezone.now())
    user_ex.save()

    data['status'] = 'SUCCESS'
    data['success'] = True
    data['message'] = 'OK'
    return JsonResponse(data)


class ForgetPwdView(View):
    def get(self, request):
        forget_form = ForgotPasswordForm()
        return render(request, "forgetpwd.html", {"forget_form": forget_form})

    def post(self, request):
        redirect_to = reverse('login')
        form = ForgotPasswordForm(request.POST)
        if form.is_valid():
            email = request.POST.get('email')
            if not UserProfile.objects.filter(email=email):
                return render(request, 'forgetpwd.html', {'msg': '此邮箱未注册'})
            print(email)

            code = request.POST.get('verification_code')
            print(code)
            identify_code = User_ex.objects.filter(valid_code=code)
            print(identify_code)
            if identify_code:
                for identify_codes in identify_code:
                    if identify_codes.email != email:
                        return render(request, 'forgetpwd.html', {'msg': '验证码不正确'})
            else:
                return render(request, 'forgetpwd.html', {'msg': '验证码不正确'})

            new_password = request.POST.get('new_password')
            print("新密码为"+new_password)
            user = UserProfile.objects.get(email=email)
            user.set_password(new_password)
            user.save()
            print("修改密码成功")

            return redirect(redirect_to)
            # return render(request, 'send_success.html')
        return render(request, "login.html")


# 修改邮箱的view:
class UpdateEmailView(LoginRequiredMixin, View):
    def post(self, request):
        email = request.POST.get("email", "")
        code = request.POST.get("code", "")

        existed_records = EmailVerifyRecord.objects.filter(
            email=email, code=code)
        if existed_records:
            user = request.user
            user.email = email
            user.save()
            return HttpResponse(
                '{"status":"success"}',
                content_type='application/json')
        else:
            return HttpResponse(
                '{"email":"验证码无效"}',
                content_type='application/json')


# 发送邮箱验证码的view:
class SendEmailCodeView(LoginRequiredMixin, View):
    def get(self, request):
        # 取出需要发送的邮件
        email = request.GET.get("email", "")
        print(email)

        # 不能是已注册的邮箱
        if UserProfile.objects.filter(email=email):
            return HttpResponse(
                '{"email":"邮箱已经存在"}',
                content_type='application/json')
        send_register_eamil(email, "update_email")

        return HttpResponse(
            '{"status":"success"}',
            content_type='application/json')


